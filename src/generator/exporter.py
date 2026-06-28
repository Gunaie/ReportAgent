"""报告导出 — Markdown → PDF (weasyprint) / Word (python-docx)。"""

import re
from pathlib import Path
from loguru import logger

import markdown


# ====== HTML 模板（weasyprint 渲染用） ======

_HTML_TEMPLATE = """<!DOCTYPE html>
<html lang="zh-CN">
<head>
<meta charset="utf-8">
<style>
  @page {{ size: A4; margin: 2cm; }}
  body {{ font-family: "Microsoft YaHei", "SimSun", sans-serif; font-size: 12pt; line-height: 1.8; color: #1a1a2e; }}
  h1 {{ font-size: 1.5em; border-bottom: 2px solid #2563eb; padding-bottom: 6px; margin-top: 0; }}
  h2 {{ font-size: 1.2em; color: #2563eb; margin-top: 1.2em; }}
  h3 {{ font-size: 1.05em; margin-top: 1em; }}
  table {{ width: 100%; border-collapse: collapse; margin: 1em 0; font-size: 10pt; }}
  th, td {{ border: 1px solid #e2e8f0; padding: 6px 10px; text-align: left; }}
  th {{ background: #f0f4f8; font-weight: 600; }}
  blockquote {{ border-left: 4px solid #2563eb; padding: 8px 16px; margin: 1em 0; color: #64748b; background: #f8fafc; }}
  hr {{ border: none; border-top: 1px solid #e2e8f0; margin: 1.5em 0; }}
  a {{ color: #2563eb; }}
  ul, ol {{ margin: 0.5em 0 0.5em 1.5em; }}
  li {{ margin: 0.3em 0; }}
  code {{ background: #f1f5f9; padding: 2px 6px; border-radius: 4px; font-size: 10pt; }}
  pre {{ background: #1a1a2e; color: #e2e8f0; padding: 12px; border-radius: 6px; overflow-x: auto; }}
  .disclaimer {{ color: #94a3b8; font-size: 10pt; margin-top: 2em; border-top: 1px solid #e2e8f0; padding-top: 1em; }}
</style>
</head>
<body>
{content}
</body>
</html>"""


def _md_to_html(md_text: str, title: str = "研究报告") -> str:
    """Markdown → 完整 HTML 页面。"""
    # 使用 Python markdown 库转换，扩展支持表格/代码/围栏代码
    md = markdown.Markdown(extensions=["tables", "fenced_code", "codehilite", "nl2br"])
    body = md.convert(md_text)
    return _HTML_TEMPLATE.format(content=body)


def _md_to_docx_text(md_text: str):
    """Markdown → python-docx 段落/表格序列。

    Yields:
        ("paragraph", text) 或 ("heading", level, text) 或 ("table", rows)
    """
    lines = md_text.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i]

        # 标题
        if line.startswith("#### "):
            yield ("heading", 4, line[5:])
        elif line.startswith("### "):
            yield ("heading", 3, line[4:])
        elif line.startswith("## "):
            yield ("heading", 2, line[3:])
        elif line.startswith("# "):
            yield ("heading", 1, line[2:])

        # 表格
        elif line.startswith("|") and line.endswith("|"):
            rows = []
            while i < len(lines) and lines[i].startswith("|"):
                row = lines[i]
                cells = [c.strip() for c in row.split("|") if c.strip()]
                if not all(c.replace("-", "").replace(":", "").replace(" ", "") == "" for c in cells):
                    rows.append(cells)
                i += 1
            if rows:
                yield ("table", rows)
            continue

        # 分割线
        elif line.strip() == "---":
            yield ("paragraph", "─" * 40)

        # 引用
        elif line.startswith("> "):
            yield ("quote", line[2:])

        # 列表项
        elif line.strip().startswith(("- ", "* ", "+ ")) or (
            len(line) > 2 and line[0].isdigit() and line[1] == "." and line[2] == " "
        ):
            text = line.strip()
            if text.startswith(("- ", "* ", "+ ")):
                text = "• " + text[2:]
            yield ("paragraph", text)

        # 空行
        elif line.strip() == "":
            yield ("paragraph", "")

        # 普通段落
        else:
            text = line
            # 链接 → 纯文本
            text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", text)
            if text.strip():
                yield ("paragraph", text)

        i += 1


# ====== 导出函数 ======


def md_to_pdf(md_text: str, output_path: Path, title: str = "研究报告") -> Path:
    """Markdown → PDF（weasyprint）。

    Args:
        md_text: Markdown 文本
        output_path: 输出 .pdf 文件路径
        title: 报告标题

    Returns:
        生成的 PDF 文件路径

    Raises:
        ImportError: weasyprint 未安装或 GTK 库缺失
    """
    try:
        from weasyprint import HTML
    except OSError as e:
        raise OSError(
            "weasyprint 需要 GTK3 运行时库。Windows: 安装 GTK3 或使用 WSL。"
            "详见: https://doc.courtbouillon.org/weasyprint/stable/first_steps.html"
        ) from e

    html_str = _md_to_html(md_text, title)
    HTML(string=html_str).write_pdf(str(output_path))
    logger.info(f"PDF 已导出: {output_path}")
    return output_path


def md_to_docx(md_text: str, output_path: Path, title: str = "研究报告") -> Path:
    """Markdown → Word (.docx)。

    Args:
        md_text: Markdown 文本
        output_path: 输出 .docx 文件路径
        title: 报告标题

    Returns:
        生成的 docx 文件路径
    """
    from docx import Document
    from docx.shared import Pt, Inches, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # 页面设置
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # 标题
    title_para = doc.add_heading(title, level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 逐块生成内容
    for item in _md_to_docx_text(md_text):
        kind = item[0]

        if kind == "heading":
            _, level, text = item
            doc.add_heading(text, level=level)
        elif kind == "paragraph":
            text = item[1]
            if text:
                para = doc.add_paragraph(text)
                para.style.font.size = Pt(11)
        elif kind == "quote":
            para = doc.add_paragraph(item[1])
            para.style.font.size = Pt(10)
            para.style.font.color.rgb = RGBColor(100, 116, 139)
        elif kind == "table":
            rows = item[1]
            if len(rows) < 2:
                continue
            table = doc.add_table(rows=len(rows), cols=len(rows[0]))
            table.style = "Light Grid Accent 1"
            for ri, row in enumerate(rows):
                for ci, cell_text in enumerate(row):
                    cell = table.cell(ri, ci)
                    cell.text = cell_text
                    if ri == 0:
                        for para in cell.paragraphs:
                            for run in para.runs:
                                run.bold = True

    # 页脚免责声明
    doc.add_paragraph("")
    disclaimer = doc.add_paragraph(
        "本报告由 AI 自动生成，基于公开信息搜索和分析。报告内容仅供研究参考，不构成投资建议。"
    )
    disclaimer.style.font.size = Pt(9)
    disclaimer.style.font.color.rgb = RGBColor(148, 163, 184)

    doc.save(str(output_path))
    logger.info(f"Word 已导出: {output_path}")
    return output_path
